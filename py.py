import json
import re
from docx import Document
from http.server import HTTPServer, SimpleHTTPRequestHandler
import webbrowser
import os

def parse_docx(file_path):
    doc = Document(file_path)
    
    # Получаем текст из всех таблиц документа
    answers = {}
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            for i in range(0, len(cells), 2):
                if i + 1 < len(cells) and cells[i] and cells[i+1]:
                    try:
                        q_num = int(cells[i])
                        answer = cells[i+1].strip()
                        if answer in ['A', 'B', 'C', 'D', 'E']:
                            answers[q_num] = answer
                    except ValueError:
                        continue

    # Получаем текст вопросов из параграфов
    text = "\n".join([para.text for para in doc.paragraphs])
    questions_text = text.split("Ключ ответов:")[0].strip()
    
    # Паттерн для вопросов
    question_pattern = re.compile(r"(\d+)\.\s(.+?)\nA\)\s(.+?)\nB\)\s(.+?)\nC\)\s(.+?)\nD\)\s(.+?)(?:\nE\)\s(.+?))?\n", re.DOTALL)
    
    questions = []
    matches = question_pattern.findall(questions_text)
    
    for match in matches:
        q_id, question, a, b, c, d, e = match if len(match) == 7 else (*match, None)
        q_id = int(q_id)
        
        options = {
            "A": a.strip(),
            "B": b.strip(),
            "C": c.strip(),
            "D": d.strip()
        }
        
        if e and e.strip():
            options["E"] = e.strip()
        
        questions.append({
            "id": q_id,
            "question": question.strip(),
            "options": options,
            "correct": answers.get(q_id, "N/A")
        })
    
    return {"questions": sorted(questions, key=lambda x: x["id"])}

def start_server():
    # Change to the directory containing our files
    os.chdir(os.path.dirname(os.path.abspath(__file__)))
    
    # Configure and start HTTP server
    server_address = ('', 8000)
    httpd = HTTPServer(server_address, SimpleHTTPRequestHandler)
    print("Server started at http://localhost:8000")
    
    # Open browser
    webbrowser.open('http://localhost:8000')
    
    # Start server
    httpd.serve_forever()

if __name__ == "__main__":
    file_path = "240 вопросов с ключом.docx"
    parsed_data = parse_docx(file_path)
    
    with open("questions.json", "w", encoding="utf-8") as json_file:
        json.dump(parsed_data, json_file, ensure_ascii=False, indent=2)
    
    print("Файл questions.json успешно создан!")
    
    # Start web server
    start_server()